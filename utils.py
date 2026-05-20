# utils.py

from __future__ import annotations

import base64
import csv
import io
import json
import re
from typing import Any, Dict, List, Optional, Tuple


def get_header(headers: List[Dict[str, Any]], name: str) -> Optional[str]:
    """Return a header value (case-insensitive) from Gmail headers list."""
    name_l = name.lower()
    for h in headers or []:
        if (h.get("name") or "").lower() == name_l:
            return h.get("value")
    return None


def _add_base64_padding(data: str) -> str:
    """Gmail uses URL-safe base64 and may omit padding."""
    return data + "=" * (-len(data) % 4)


def decode_body_data_to_bytes(data: str) -> bytes:
    """Decode Gmail's URL-safe base64 encoded data to bytes."""
    if not data:
        return b""
    try:
        return base64.urlsafe_b64decode(_add_base64_padding(data))
    except Exception:
        return b""


def _decode_body_data(data: str) -> str:
    """Decode Gmail's URL-safe base64 encoded data to text."""
    return decode_body_data_to_bytes(data).decode("utf-8", errors="replace")


def extract_plain_text_from_payload(payload: Dict[str, Any]) -> str:
    """
    Walk Gmail message payload recursively and return first text/plain content found.
    """
    if not payload:
        return ""

    mime = payload.get("mimeType", "")

    # Direct text/plain
    if mime == "text/plain":
        body = payload.get("body", {}) or {}
        data = body.get("data")
        if data:
            return _decode_body_data(data)

    # Multipart: iterate parts
    for part in payload.get("parts", []) or []:
        txt = extract_plain_text_from_payload(part)
        if txt:
            return txt

    return ""


def sanitize_filename(filename: str, fallback: str = "attachment") -> str:
    """Return a safe filename for Content-Disposition/download names."""
    filename = (filename or "").strip().replace("\\", "_").replace("/", "_")
    filename = re.sub(r"[^A-Za-z0-9._()\- ]+", "_", filename).strip(" .")
    return filename or fallback


def _walk_payload_parts(payload: Dict[str, Any]):
    """Yield every payload part, including nested multipart children."""
    if not payload:
        return
    yield payload
    for part in payload.get("parts", []) or []:
        yield from _walk_payload_parts(part)


def list_attachments_from_payload(payload: Dict[str, Any]) -> List[Dict[str, Any]]:
    """
    Return metadata for visible Gmail attachments.

    Gmail stores downloadable file attachments as MIME parts with a filename.
    The file bytes are either in body.data for very small parts or behind
    body.attachmentId for normal attachments.
    """
    attachments: List[Dict[str, Any]] = []

    for part in _walk_payload_parts(payload):
        filename = (part.get("filename") or "").strip()
        body = part.get("body", {}) or {}
        attachment_id = body.get("attachmentId") or ""
        inline_data = bool(body.get("data"))

        # Gmail may include nameless multipart/plain/html parts; do not show those.
        if not filename:
            continue

        attachments.append(
            {
                "index": len(attachments),
                "filename": sanitize_filename(filename),
                "original_filename": filename,
                "mimeType": part.get("mimeType") or "application/octet-stream",
                "size": int(body.get("size") or 0),
                "attachmentId": attachment_id,
                "partId": part.get("partId") or "",
                "hasInlineData": inline_data,
                "canPreview": can_preview_mime(part.get("mimeType") or ""),
            }
        )

    return attachments


def find_attachment_part(
    payload: Dict[str, Any],
    *,
    attachment_id: str = "",
    filename: str = "",
    part_id: str = "",
    index: Optional[int] = None,
) -> Optional[Dict[str, Any]]:
    """Find an attachment MIME part by ID, part id, filename, or displayed index."""
    matches = []
    safe_filename = sanitize_filename(filename) if filename else ""

    for part in _walk_payload_parts(payload):
        part_filename = (part.get("filename") or "").strip()
        if not part_filename:
            continue

        body = part.get("body", {}) or {}
        matches.append(part)

        if attachment_id and body.get("attachmentId") == attachment_id:
            return part
        if part_id and part.get("partId") == part_id:
            return part
        if safe_filename and sanitize_filename(part_filename) == safe_filename:
            return part

    if index is not None and 0 <= index < len(matches):
        return matches[index]

    return None


def can_preview_mime(mime_type: str) -> bool:
    """Whether the browser can usually preview this attachment inline."""
    mime_type = (mime_type or "").lower()
    return (
        mime_type.startswith("image/")
        or mime_type.startswith("text/")
        or mime_type in {
            "application/pdf",
            "application/json",
            "application/xml",
            "text/html",
            "text/csv",
        }
    )


def _looks_text_like_mime(mime_type: str, filename: str = "") -> bool:
    """Return True when bytes are likely safe to decode directly as text."""
    mime_type = (mime_type or "").lower()
    filename = (filename or "").lower()
    text_extensions = (
        ".txt", ".csv", ".json", ".xml", ".html", ".htm", ".md",
        ".log", ".ics", ".vcf", ".yaml", ".yml", ".py", ".js",
        ".ts", ".css", ".java", ".c", ".cpp", ".sql",
    )
    return (
        mime_type.startswith("text/")
        or mime_type in {
            "application/json",
            "application/xml",
            "application/xhtml+xml",
            "application/csv",
            "text/csv",
        }
        or filename.endswith(text_extensions)
    )


def _decode_possible_text(data: bytes) -> str:
    """Decode bytes using common encodings without failing hard."""
    if not data:
        return ""
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return data.decode(encoding, errors="replace")
        except Exception:
            continue
    return data.decode("utf-8", errors="replace")


def _clean_extracted_text(text: str, max_chars: int) -> Tuple[str, str]:
    """Normalize whitespace and apply a character cap for summarization."""
    text = (text or "").replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    if len(text) > max_chars:
        return text[:max_chars].rstrip(), f"Text was truncated to {max_chars:,} characters for summarization."
    return text, ""


def extract_text_from_attachment_bytes(
    data: bytes,
    *,
    filename: str = "",
    mime_type: str = "",
    max_chars: int = 20000,
) -> Tuple[str, str]:
    """
    Extract readable text from common attachment formats for AI summarization.

    Supported well: text/CSV/JSON/XML/HTML, PDF, DOCX, XLSX.
    Unsupported binary files return an empty text string plus a helpful note.
    """
    filename_l = (filename or "").lower()
    mime_l = (mime_type or "").lower()

    if not data:
        return "", "Attachment file data was empty."

    try:
        if _looks_text_like_mime(mime_l, filename_l):
            text = _decode_possible_text(data)

            # Pretty-print JSON when possible to improve summaries.
            if mime_l == "application/json" or filename_l.endswith(".json"):
                try:
                    text = json.dumps(json.loads(text), indent=2, ensure_ascii=False)
                except Exception:
                    pass

            return _clean_extracted_text(text, max_chars)

        if mime_l == "application/pdf" or filename_l.endswith(".pdf"):
            try:
                from pypdf import PdfReader
            except Exception:
                return "", "PDF text extraction needs pypdf. Run: pip install pypdf"

            reader = PdfReader(io.BytesIO(data))
            page_texts = []
            for page_num, page in enumerate(reader.pages, start=1):
                try:
                    page_text = page.extract_text() or ""
                except Exception:
                    page_text = ""
                if page_text.strip():
                    page_texts.append(f"[Page {page_num}]\n{page_text.strip()}")
            if not page_texts:
                return "", "No selectable text was found in this PDF. It may be scanned or image-only."
            return _clean_extracted_text("\n\n".join(page_texts), max_chars)

        if (
            mime_l == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            or filename_l.endswith(".docx")
        ):
            try:
                from docx import Document
            except Exception:
                return "", "DOCX text extraction needs python-docx. Run: pip install python-docx"

            doc = Document(io.BytesIO(data))
            chunks = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        chunks.append(" | ".join(cells))
            if not chunks:
                return "", "No readable text was found in this DOCX file."
            return _clean_extracted_text("\n".join(chunks), max_chars)

        if (
            mime_l == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            or filename_l.endswith(".xlsx")
        ):
            try:
                from openpyxl import load_workbook
            except Exception:
                return "", "XLSX text extraction needs openpyxl. Run: pip install openpyxl"

            workbook = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
            rows_out = []
            max_rows_per_sheet = 80
            max_cols = 12
            for sheet in workbook.worksheets[:5]:
                rows_out.append(f"[Sheet: {sheet.title}]")
                for row_index, row in enumerate(sheet.iter_rows(values_only=True), start=1):
                    if row_index > max_rows_per_sheet:
                        rows_out.append("...(sheet truncated)...")
                        break
                    values = [str(value).strip() if value is not None else "" for value in row[:max_cols]]
                    if any(values):
                        rows_out.append(" | ".join(values))
            text = "\n".join(rows_out).strip()
            if not text:
                return "", "No readable values were found in this XLSX file."
            return _clean_extracted_text(text, max_chars)

    except Exception as exc:
        return "", f"Could not extract attachment text: {exc}"

    return "", "This attachment type is not text-extractable by the app yet. Preview or download it instead."
